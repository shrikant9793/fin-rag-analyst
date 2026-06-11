"""
src/rag_pipeline.py
===================
Hybrid RAG Pipeline — Qdrant (dense + sparse) with RRF reranking.

Architecture:
  1. Dense search  → BGE-M3 embeddings (sentence-transformers, local)
  2. Sparse search → BM42 embeddings  (fastembed, local)
  3. RRF fusion   → Reciprocal Rank Fusion merges both result lists
  4. Returns top-N chunks as LangChain Document objects

Public API:
    pipeline = RAGPipeline()
    pipeline.ingest_documents(file_paths, metadata)   # Day-1 ingestion
    docs = pipeline.retrieve(query, top_n=5)          # Day-2+ retrieval
"""

from __future__ import annotations

import hashlib
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import numpy as np
from langchain_core.documents import Document
from loguru import logger
from qdrant_client import QdrantClient, models
from qdrant_client.http.exceptions import UnexpectedResponse
from qdrant_client.models import (
    Distance,
    FieldCondition,
    Filter,
    MatchValue,
    PointStruct,
    SparseVector,
    SparseVectorParams,
    VectorParams,
)
from sentence_transformers import SentenceTransformer
from tenacity import retry, stop_after_attempt, wait_exponential

from src.config import get_settings

# ---------------------------------------------------------------------------
# Text splitter (LangChain)
# ---------------------------------------------------------------------------
from langchain_text_splitters import RecursiveCharacterTextSplitter


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------
@dataclass
class RetrievedChunk:
    """Single retrieved document chunk with score and metadata."""

    content: str
    score: float
    metadata: dict[str, Any] = field(default_factory=dict)
    chunk_id: str = ""

    def to_langchain_doc(self) -> Document:
        return Document(page_content=self.content, metadata={**self.metadata, "score": self.score})


# ---------------------------------------------------------------------------
# RAG Pipeline
# ---------------------------------------------------------------------------
class RAGPipeline:
    """
    Manages embedding, ingestion, and hybrid retrieval against Qdrant.

    Example
    -------
    >>> pipeline = RAGPipeline()
    >>> pipeline.ensure_collection()
    >>> pipeline.ingest_documents(
    ...     file_paths=[Path("data/raw/AAPL_10K_2024.pdf")],
    ...     doc_metadata={"ticker": "AAPL", "doc_type": "sec_filing", "filing_date": "2024-10-31"},
    ... )
    >>> chunks = pipeline.retrieve("What is Apple's gross margin for FY2024?")
    """

    def __init__(self) -> None:
        self.settings   = get_settings()
        self.qdrant_cfg = self.settings.qdrant_config
        self.ing_cfg    = self.settings.ingestion_config
        self.emb_dense  = self.settings.embedding_dense
        self.emb_sparse = self.settings.embedding_sparse

        self._client: QdrantClient | None        = None
        self._dense_model: SentenceTransformer | None = None
        self._sparse_model: Any | None           = None   # fastembed TextEmbedding

        logger.info("RAGPipeline initialised")

    # -----------------------------------------------------------------------
    # Lazy initialisation helpers
    # -----------------------------------------------------------------------
    @property
    def client(self) -> QdrantClient:
        if self._client is None:
            self._client = self._connect()
        return self._client

    @property
    def dense_model(self) -> SentenceTransformer:
        if self._dense_model is None:
            logger.info(f"Loading dense embedding model: {self.emb_dense['model']}")
            self._dense_model = SentenceTransformer(
                self.emb_dense["model"],
                device=self.emb_dense["device"],
            )
        return self._dense_model

    @property
    def sparse_model(self):
        if self._sparse_model is None:
            logger.info(f"Loading sparse embedding model: {self.emb_sparse['model']}")
            from fastembed import SparseTextEmbedding
            self._sparse_model = SparseTextEmbedding(model_name=self.emb_sparse["model"])
        return self._sparse_model

    # -----------------------------------------------------------------------
    # Qdrant Connection
    # -----------------------------------------------------------------------
    @retry(stop=stop_after_attempt(5), wait=wait_exponential(multiplier=1, min=2, max=10))
    def _connect(self) -> QdrantClient:
        url     = self.qdrant_cfg["url"]
        api_key = self.qdrant_cfg.get("api_key") or None
        logger.info(f"Connecting to Qdrant at {url}")
        client = QdrantClient(url=url, api_key=api_key, timeout=30)
        # Ping
        client.get_collections()
        logger.success(f"Qdrant connected: {url}")
        return client

    # -----------------------------------------------------------------------
    # Collection Management
    # -----------------------------------------------------------------------
    def ensure_collection(self, recreate: bool = False) -> None:
        """
        Create the Qdrant collection with both dense and sparse vector configs.
        Skips creation if already exists (unless recreate=True).
        """
        name        = self.qdrant_cfg["collection_name"]
        vector_size = self.qdrant_cfg["vector_size"]
        distance    = Distance[self.qdrant_cfg["distance"].upper()]

        existing = {c.name for c in self.client.get_collections().collections}

        if name in existing:
            if not recreate:
                logger.info(f"Collection '{name}' already exists — skipping creation")
                return
            logger.warning(f"Recreating collection '{name}'")
            self.client.delete_collection(name)

        self.client.create_collection(
            collection_name=name,
            vectors_config={
                "dense": VectorParams(size=vector_size, distance=distance),
            },
            sparse_vectors_config={
                "sparse": SparseVectorParams(),
            },
        )

        # Payload indexes for fast metadata filtering
        for field_name in ["ticker", "doc_type", "filing_date"]:
            self.client.create_payload_index(
                collection_name=name,
                field_name=field_name,
                field_schema="keyword",
            )

        logger.success(f"Collection '{name}' created with dense+sparse vectors")

    # -----------------------------------------------------------------------
    # Embedding helpers
    # -----------------------------------------------------------------------
    def _embed_dense(self, texts: list[str]) -> list[list[float]]:
        """Batch embed texts with BGE-M3."""
        vectors = self.dense_model.encode(
            texts,
            batch_size=self.emb_dense["batch_size"],
            normalize_embeddings=self.emb_dense["normalize"],
            show_progress_bar=len(texts) > 20,
        )
        return vectors.tolist()

    def _embed_sparse(self, texts: list[str]) -> list[SparseVector]:
        """Batch embed texts with BM42 (fastembed) → SparseVector list."""
        results = list(self.sparse_model.embed(texts))
        sparse_vectors = []
        for result in results:
            indices = result.indices.tolist()
            values  = result.values.tolist()
            sparse_vectors.append(SparseVector(indices=indices, values=values))
        return sparse_vectors

    # -----------------------------------------------------------------------
    # Document Loading & Chunking
    # -----------------------------------------------------------------------
    def _load_file(self, path: Path) -> str:
        """Load raw text from PDF, TXT, or DOCX."""
        suffix = path.suffix.lower()

        if suffix == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")

        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        if suffix == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs)

        raise ValueError(f"Unsupported file type: {suffix}")

    def _chunk_text(self, text: str) -> list[str]:
        """Split text into overlapping chunks."""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.ing_cfg["chunk_size"],
            chunk_overlap=self.ing_cfg["chunk_overlap"],
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        return splitter.split_text(text)

    # -----------------------------------------------------------------------
    # Ingestion
    # -----------------------------------------------------------------------
    def ingest_documents(
        self,
        file_paths: list[Path],
        doc_metadata: dict[str, str],
        batch_size: int = 64,
    ) -> int:
        """
        Load, chunk, embed, and upsert documents into Qdrant.

        Args:
            file_paths:   List of paths to financial documents.
            doc_metadata: Shared metadata applied to all chunks
                          (ticker, doc_type, filing_date, company_name).
            batch_size:   Upsert batch size.

        Returns:
            Total number of chunks ingested.
        """
        self.ensure_collection()
        collection = self.qdrant_cfg["collection_name"]
        total_chunks = 0

        for file_path in file_paths:
            path = Path(file_path)
            if not path.exists():
                logger.warning(f"File not found, skipping: {path}")
                continue

            logger.info(f"Ingesting: {path.name}")
            t0   = time.perf_counter()
            text = self._load_file(path)

            if not text.strip():
                logger.warning(f"Empty content extracted from {path.name}")
                continue

            chunks = self._chunk_text(text)
            logger.info(f"  → {len(chunks)} chunks from {path.name}")

            # Process in batches
            for i in range(0, len(chunks), batch_size):
                batch_texts = chunks[i : i + batch_size]

                dense_vecs  = self._embed_dense(batch_texts)
                sparse_vecs = self._embed_sparse(batch_texts)

                points = []
                for j, (chunk_text, dense_v, sparse_v) in enumerate(
                    zip(batch_texts, dense_vecs, sparse_vecs)
                ):
                    chunk_id = hashlib.md5(
                        f"{path.name}_{i + j}_{chunk_text[:50]}".encode()
                    ).hexdigest()

                    points.append(
                        PointStruct(
                            id=str(uuid.uuid4()),
                            vector={"dense": dense_v, "sparse": sparse_v},
                            payload={
                                "content": chunk_text,
                                "chunk_id": chunk_id,
                                "source_file": path.name,
                                "chunk_index": i + j,
                                **doc_metadata,
                            },
                        )
                    )

                self.client.upsert(collection_name=collection, points=points)
                logger.debug(f"  Upserted batch {i // batch_size + 1} ({len(points)} points)")

            total_chunks += len(chunks)
            elapsed = time.perf_counter() - t0
            logger.success(f"  ✓ {path.name} ingested in {elapsed:.2f}s ({len(chunks)} chunks)")

        logger.success(f"Ingestion complete — {total_chunks} total chunks in '{collection}'")
        return total_chunks

    # -----------------------------------------------------------------------
    # Retrieval — Hybrid Search with RRF
    # -----------------------------------------------------------------------
    def retrieve(
        self,
        query: str,
        top_n: int | None = None,
        filters: dict[str, str] | None = None,
    ) -> list[RetrievedChunk]:
        """
        Run hybrid search (dense + sparse) and fuse with RRF.

        Args:
            query:   Natural language financial query.
            top_n:   Number of final chunks to return (default from config).
            filters: Optional payload filters e.g. {"ticker": "AAPL"}.

        Returns:
            List of RetrievedChunk sorted by RRF score (descending).
        """
        ret_cfg    = self.qdrant_cfg["retrieval"]
        top_k      = ret_cfg["top_k"]
        rrf_k      = ret_cfg["rrf_k"]
        final_n    = top_n or ret_cfg["rerank_top_n"]
        collection = self.qdrant_cfg["collection_name"]

        logger.info(f"Hybrid search | query='{query[:80]}...' | top_k={top_k}")

        # Build Qdrant filter if metadata filters provided
        qdrant_filter = None
        if filters:
            conditions = [
                FieldCondition(key=k, match=MatchValue(value=v))
                for k, v in filters.items()
            ]
            qdrant_filter = Filter(must=conditions)

        # --- Dense Search ---
        query_dense_vec = self._embed_dense([query])[0]
        dense_results = self.client.search(
            collection_name=collection,
            query_vector=models.NamedVector(name="dense", vector=query_dense_vec),
            limit=top_k,
            query_filter=qdrant_filter,
            with_payload=True,
        )

        # --- Sparse Search ---
        query_sparse_vec = self._embed_sparse([query])[0]
        sparse_results = self.client.search(
            collection_name=collection,
            query_vector=models.NamedSparseVector(name="sparse", vector=query_sparse_vec),
            limit=top_k,
            query_filter=qdrant_filter,
            with_payload=True,
        )

        # --- RRF Fusion ---
        fused = self._rrf_fuse(dense_results, sparse_results, k=rrf_k)

        # Build RetrievedChunk objects
        chunks = []
        for point_id, rrf_score in fused[:final_n]:
            # Find payload from either result list
            payload = {}
            for result in dense_results + sparse_results:
                if result.id == point_id:
                    payload = result.payload or {}
                    break

            chunks.append(
                RetrievedChunk(
                    content=payload.get("content", ""),
                    score=rrf_score,
                    metadata={k: v for k, v in payload.items() if k != "content"},
                    chunk_id=payload.get("chunk_id", str(point_id)),
                )
            )

        logger.info(f"Retrieved {len(chunks)} chunks after RRF fusion")
        return chunks

    @staticmethod
    def _rrf_fuse(
        dense_results: list,
        sparse_results: list,
        k: int = 60,
    ) -> list[tuple[str, float]]:
        """
        Reciprocal Rank Fusion of two ranked lists.

        Score(doc) = Σ 1 / (k + rank_i)   for each list i that contains doc.
        """
        scores: dict[str, float] = {}

        for rank, result in enumerate(dense_results, start=1):
            doc_id = str(result.id)
            scores[doc_id] = scores.get(doc_id, 0.0) + 1.0 / (k + rank)

        for rank, result in enumerate(sparse_results, start=1):
            doc_id = str(result.id)
            scores[doc_id] = scores.get(doc_id, 0.0) + 1.0 / (k + rank)

        # Sort descending by RRF score
        return sorted(scores.items(), key=lambda x: x[1], reverse=True)

    # -----------------------------------------------------------------------
    # Utility
    # -----------------------------------------------------------------------
    def collection_stats(self) -> dict:
        """Return count and vector config of the active collection."""
        name = self.qdrant_cfg["collection_name"]
        info = self.client.get_collection(name)
        return {
            "collection": name,
            "points_count": info.points_count,
            "vectors_count": info.vectors_count,
            "status": info.status,
        }

    def delete_collection(self) -> None:
        """Drop the collection — use with caution."""
        name = self.qdrant_cfg["collection_name"]
        self.client.delete_collection(name)
        logger.warning(f"Collection '{name}' deleted")